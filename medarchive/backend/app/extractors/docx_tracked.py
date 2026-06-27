"""DOCX extractor with tracked-changes resolution (brief §7, §20).

python-docx ignores tracked changes, so we walk the XML: keep w:ins (accepted
insertions), drop w:del/w:delText (deletions), to recover the FINAL text. Then
extract tables into service/price rows.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.extractors.base import (
    ExtractedRow,
    ExtractionResult,
    classify_columns,
    looks_like_header,
    register,
)
from app.models.enums import FileFormat

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def cell_final_text(cell) -> str:
    """Resolve tracked changes inside a table cell to final accepted text."""
    parts: list[str] = []
    tc = cell._tc
    for node in tc.iter():
        tag = node.tag
        if tag == f"{_W}t":
            # Skip text that lives inside a w:del (deleted) ancestor.
            if _has_ancestor(node, f"{_W}del"):
                continue
            parts.append(node.text or "")
        # w:delText is never included (it's the deleted content).
    return "".join(parts).strip()


def _has_ancestor(node, tag: str) -> bool:
    parent = node.getparent()
    while parent is not None:
        if parent.tag == tag:
            return True
        parent = parent.getparent()
    return False


@register(FileFormat.docx)
class DocxTrackedExtractor:
    def extract(self, path: Path) -> ExtractionResult:
        document = Document(path)
        rows: list[ExtractedRow] = []
        text_lines: list[str] = []
        notes: list[str] = []

        # Title/paragraph context (first non-empty paragraph) as a weak hint.
        title = ""
        for p in document.paragraphs:
            if p.text.strip():
                title = p.text.strip()
                break
        if title:
            text_lines.append(f"# {title}")

        for ti, table in enumerate(document.tables):
            grid = [[cell_final_text(c) for c in row.cells] for row in table.rows]
            if not grid:
                continue
            header_idx = next(
                (i for i, r in enumerate(grid[:5]) if looks_like_header(r)), None
            )
            if header_idx is None:
                header_idx = 0
            cols = classify_columns(grid[header_idx])
            if cols["name"] is None:
                cols["name"] = 0
            notes.append(f"table{ti}: header@{header_idx} cols={cols}")

            for r in grid[header_idx + 1:]:
                line = " | ".join(r)
                text_lines.append(line)
                name = self._at(r, cols["name"])
                if not name:
                    continue
                rows.append(
                    ExtractedRow(
                        service_name_raw=name,
                        price_resident=self._at(r, cols["resident"]),
                        price_nonresident=self._at(r, cols["nonresident"]),
                        source_fragment=line[:200],
                    )
                )
        return ExtractionResult(rows=rows, raw_text="\n".join(text_lines), log="; ".join(notes))

    @staticmethod
    def _at(row: list[str], idx: int | None) -> str | None:
        if idx is None or idx >= len(row):
            return None
        return (row[idx] or "").strip() or None
