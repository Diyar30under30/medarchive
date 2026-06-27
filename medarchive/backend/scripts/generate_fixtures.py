"""Generate `data/incoming/sample_archive.zip` — synthetic partner price lists
across all four formats with deliberate messiness and anomalies (brief §5).

Also writes `data/incoming/fixtures_manifest.json` (ground truth) for tests.

Run: python -m scripts.generate_fixtures
"""
from __future__ import annotations

import io
import json
import zipfile
from datetime import date
from pathlib import Path

from app.config import settings
from scripts._fixture_data import CLINICS, RAW_TO_CANONICAL

# ── price assignment (deterministic) ──────────────────────────────────────────
def _price_for(raw: str) -> int:
    base = 1500 + (abs(hash(raw)) % 18) * 500  # 1500..10000, stable per name
    return base


def _rows_for_clinic(seed: int, n: int) -> list[dict]:
    """Pick n raw services for a clinic, with resident/non-resident KZT prices."""
    chosen = [RAW_TO_CANONICAL[(seed + i) % len(RAW_TO_CANONICAL)] for i in range(n)]
    rows = []
    for raw, canon, hint in chosen:
        res = _price_for(raw)
        rows.append(
            {
                "raw": raw,
                "canonical": canon,
                "specialty_hint": hint,
                "resident": res,
                "nonresident": int(res * 1.4),
                "currency": "KZT",
            }
        )
    return rows


_FONT_CANDIDATES = [
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/segoeui.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
]


def _cyrillic_fontfile() -> str | None:
    """A TTF with Cyrillic glyphs; PyMuPDF's base fonts are Latin-only, so
    inserting Cyrillic without this produces unreadable (unextractable) text."""
    from os.path import exists

    for p in _FONT_CANDIDATES:
        if exists(p):
            return p
    return None


_FONT = _cyrillic_fontfile()


def _put(page, point, text, fontsize):
    if _FONT:
        page.insert_text(point, text, fontsize=fontsize, fontname="cyr", fontfile=_FONT)
    else:
        page.insert_text(point, text, fontsize=fontsize)


# ── format writers ────────────────────────────────────────────────────────────
def _meta_lines(clinic: dict) -> list[str]:
    lines = [
        f"Клиника: {clinic['name']}",
        f"Город: {clinic['city']}",
        f"Адрес: {clinic['address']}",
        f"Тел: {clinic['phone']}",
    ]
    if clinic.get("bin"):
        lines.append(f"БИН: {clinic['bin']}")
    return lines


def write_xlsx(clinic: dict, rows: list[dict], *, header_offset: int = 0,
               multi_sheet: bool = False) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Прайс"
    # Clinic metadata preamble — also pushes the header off row 1 (brief: detect).
    for line in _meta_lines(clinic):
        ws.append([line])
    for _ in range(header_offset):
        ws.append([""])
    ws.append(["Наименование услуги", "Цена резидент", "Цена нерезидент"])
    for r in rows:
        ws.append([r["raw"], r["resident"], r["nonresident"]])
    if multi_sheet:
        ws2 = wb.create_sheet("Доп. услуги")
        ws2.append(["info"])  # junk preamble
        ws2.append(["Услуга", "Стоимость"])
        for r in rows[: max(1, len(rows) // 2)]:
            ws2.append([r["raw"], r["resident"]])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def write_pdf_text(clinic: dict, rows: list[dict]) -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    y = 50
    _put(page, (50, y), f"{clinic['name']} — прайс-лист", 14)
    y += 18
    for line in _meta_lines(clinic):
        _put(page, (50, y), line, 9)
        y += 13
    y += 6
    _put(page, (50, y), "Услуга / Резидент / Нерезидент", 10)
    y += 20
    for r in rows:
        _put(page, (50, y), f"{r['raw']}    {r['resident']}    {r['nonresident']}", 10)
        y += 16
        if y > 760:
            page = doc.new_page()
            y = 60
    data = doc.tobytes()
    doc.close()
    return data


def write_pdf_scan(clinic: dict, rows: list[dict]) -> bytes:
    """Render text to an image and embed as an image-only PDF (no text layer)."""
    import fitz

    # First render a text page, rasterize it, then place the image on a blank PDF.
    tmp = fitz.open()
    p = tmp.new_page()
    y = 50
    _put(p, (50, y), f"{clinic['name']} (скан)", 16)
    y += 24
    for line in _meta_lines(clinic):
        _put(p, (50, y), line, 11)
        y += 18
    y += 6
    for r in rows:
        _put(p, (50, y), f"{r['raw']}   {r['resident']}   {r['nonresident']}", 13)
        y += 22
    pix = p.get_pixmap(dpi=150)
    img_bytes = pix.tobytes("png")
    tmp.close()

    out = fitz.open()
    page = out.new_page(width=pix.width * 72 / 150, height=pix.height * 72 / 150)
    page.insert_image(page.rect, stream=img_bytes)
    data = out.tobytes()
    out.close()
    return data


def write_docx_tracked(clinic: dict, rows: list[dict]) -> bytes:
    """DOCX whose price table contains tracked insertions/deletions. The final
    (accepted) text is what extraction must recover."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = Document()
    document.add_heading(clinic["name"], level=1)
    for line in _meta_lines(clinic):
        document.add_paragraph(line)
    table = document.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Услуга"
    hdr[1].text = "Резидент"
    hdr[2].text = "Нерезидент"

    for i, r in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = r["raw"]
        cells[2].text = str(r["nonresident"])
        # Make the resident price a tracked insertion (w:ins) for ~half the rows,
        # and add a tracked-deleted junk run (w:del) that must be dropped.
        para = cells[1].paragraphs[0]
        if i % 2 == 0:
            _add_tracked_ins(para, str(r["resident"]))
            _add_tracked_del(para, "999999")  # deleted → must NOT appear
        else:
            para.add_run(str(r["resident"]))
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _add_tracked_ins(paragraph, text: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "editor")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    ins.append(run)
    paragraph._p.append(ins)


def _add_tracked_del(paragraph, text: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "2")
    dele.set(qn("w:author"), "editor")
    run = OxmlElement("w:r")
    delText = OxmlElement("w:delText")
    delText.text = text
    run.append(delText)
    dele.append(run)
    paragraph._p.append(dele)


# ── archive plan ──────────────────────────────────────────────────────────────
def build_archive() -> tuple[bytes, list[dict]]:
    """Returns (zip_bytes, manifest)."""
    files: list[tuple[str, bytes]] = []
    manifest: list[dict] = []

    def add(name: str, data: bytes, clinic: dict, fmt: str, eff: date,
            rows: list[dict], notes: list[str] | None = None):
        files.append((name, data))
        manifest.append(
            {
                "file_name": name,
                "format": fmt,
                "partner": clinic["name"],
                "city": clinic["city"],
                "bin": clinic["bin"],
                "effective_date": eff.isoformat(),
                "rows": rows,
                "notes": notes or [],
            }
        )

    c = CLINICS
    # 1. XLSX, header not row 1, multi-sheet
    r1 = _rows_for_clinic(0, 8)
    add("Сункар_прайс_2024-01-15.xlsx", write_xlsx(c[0], r1, header_offset=2, multi_sheet=True),
        c[0], "xlsx", date(2024, 1, 15), r1, ["header_not_row1", "multi_sheet"])

    # 2. text PDF
    r2 = _rows_for_clinic(5, 7)
    add("Сеним_2024-02-01.pdf", write_pdf_text(c[1], r2),
        c[1], "pdf", date(2024, 2, 1), r2)

    # 3. DOCX with tracked changes
    r3 = _rows_for_clinic(9, 6)
    add("ДиагностикаПлюс_2024-02-10.docx", write_docx_tracked(c[2], r3),
        c[2], "docx", date(2024, 2, 10), r3, ["tracked_changes"])

    # 4. scan PDF (OCR)
    r4 = _rows_for_clinic(2, 6)
    add("Инвиво_scan_2024-03-01.pdf", write_pdf_scan(c[3], r4),
        c[3], "scan_pdf", date(2024, 3, 1), r4, ["ocr_required"])

    # 5. XLSX simple
    r5 = _rows_for_clinic(12, 7)
    add("Поликлиника5_2024-03-05.xlsx", write_xlsx(c[4], r5),
        c[4], "xlsx", date(2024, 3, 5), r5)

    # 6a/6b. MediLab two dated files → versioning + >50% price change anomaly
    r6a = _rows_for_clinic(3, 6)
    add("MediLab_2024-01-20.pdf", write_pdf_text(c[5], r6a),
        c[5], "pdf", date(2024, 1, 20), r6a)
    r6b = [dict(x) for x in r6a]
    r6b[0]["resident"] = int(r6b[0]["resident"] * 2.2)  # >50% jump
    r6b[0]["nonresident"] = int(r6b[0]["nonresident"] * 2.2)
    add("MediLab_2024-04-20.pdf", write_pdf_text(c[5], r6b),
        c[5], "pdf", date(2024, 4, 20), r6b, ["price_jump_gt_50pct", "versioning"])

    # 7. DOCX Kazakh
    r7 = _rows_for_clinic(24, 5)
    add("Жулдыз_2024-02-15.docx", write_docx_tracked(c[6], r7),
        c[6], "docx", date(2024, 2, 15), r7, ["kazakh"])

    # 8. XLSX with the anomaly bundle
    r8 = _rows_for_clinic(6, 6)
    r8[0]["nonresident"] = r8[0]["resident"] - 500           # non-resident < resident
    r8[1]["currency"] = "USD"; r8[1]["resident"] = 50; r8[1]["nonresident"] = 70  # currency
    r8[2]["currency"] = "RUB"; r8[2]["resident"] = 3000; r8[2]["nonresident"] = 4200
    r8.append(dict(r8[3]))                                     # duplicate row
    r8.append({"raw": "", "canonical": None, "specialty_hint": None,
               "resident": 1000, "nonresident": 1400, "currency": "KZT"})  # empty name
    add("АлтынДари_2024-03-10.xlsx", write_xlsx(c[7], r8),
        c[7], "xlsx", date(2030, 1, 1), r8,
        ["nonresident_lt_resident", "currency_usd_rub", "duplicate_row",
         "empty_service_name", "future_effective_date"])

    # 9. deliberately broken file (fault tolerance) — not in manifest rows
    files.append(("broken_файл.pdf", b"%PDF-1.4 this is not a real pdf \x00\xff corrupt"))
    manifest.append({"file_name": "broken_файл.pdf", "format": "pdf",
                     "partner": None, "effective_date": None, "rows": [],
                     "notes": ["deliberately_broken"]})

    # zip it
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files:
            zf.writestr(name, data)
    return buf.getvalue(), manifest


def main() -> Path:
    zip_bytes, manifest = build_archive()
    out_dir = settings.incoming_path
    out_dir.mkdir(parents=True, exist_ok=True)
    zip_path = out_dir / "sample_archive.zip"
    zip_path.write_bytes(zip_bytes)
    (out_dir / "fixtures_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    n_rows = sum(len(m["rows"]) for m in manifest)
    print(f"Wrote {zip_path} ({len(manifest)} files, {n_rows} ground-truth rows)")
    print(f"Manifest → {out_dir / 'fixtures_manifest.json'}")
    return zip_path


if __name__ == "__main__":
    main()
